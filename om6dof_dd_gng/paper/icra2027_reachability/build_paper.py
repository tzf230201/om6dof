from __future__ import annotations

import csv
import math
import statistics
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "source_snapshot"
ASSETS = ROOT / "assets"
DATA = ROOT / "data"
ANALYSIS = DATA / "analysis_guarded_heldout_20260823"
QA = ROOT / "qa_render_heldout"
OUTPUT = ROOT / "Coverage_Connectivity_Reachability_Heldout_Draft.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
MUTED = "666666"
WHITE = "FFFFFF"
BLACK = "000000"


def pc(value):
    if isinstance(value, str) and len(value) == 6 and not value.startswith("#"):
        return f"#{value}"
    return value


def font(size: int, bold: bool = False):
    name = "arialbd.ttf" if bold else "arial.ttf"
    try:
        return ImageFont.truetype(name, size)
    except OSError:
        return ImageFont.load_default()


def rounded_box(draw, xy, fill, outline, title, lines, title_color=INK):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=22, fill=pc(fill), outline=pc(outline), width=4)
    draw.text((x1 + 24, y1 + 20), title, font=font(28, True), fill=pc(title_color))
    y = y1 + 63
    for line in lines:
        draw.text((x1 + 24, y), line, font=font(21), fill=pc(BLACK))
        y += 31


def arrow(draw, start, end, color=INK, width=7):
    draw.line([start, end], fill=pc(color), width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 20
    spread = 0.55
    p1 = (
        end[0] - length * math.cos(angle - spread),
        end[1] - length * math.sin(angle - spread),
    )
    p2 = (
        end[0] - length * math.cos(angle + spread),
        end[1] - length * math.sin(angle + spread),
    )
    draw.polygon([end, p1, p2], fill=pc(color))


def build_architecture_figure(path: Path):
    image = Image.new("RGB", (1800, 1040), pc(WHITE))
    draw = ImageDraw.Draw(image)
    draw.text((70, 35), "Dual-topology reachability planning pipeline", font=font(38, True), fill=pc(INK))
    draw.text(
        (70, 88),
        "Environment topology supplies targets and obstacles; robot topology retains executable joint configurations.",
        font=font(23),
        fill=pc(MUTED),
    )

    rounded_box(draw, (70, 180, 390, 420), "EAF3F8", BLUE, "Wrist RGB-D", ["D405 depth + color", "tf2 -> world frame", "semantic detections"])
    rounded_box(draw, (510, 160, 920, 440), "EAF3F8", BLUE, "Environment graph G_E", ["DD-GNG nodes", "labelled node = target", "unlabelled nodes/edges", "= collision geometry"])
    rounded_box(draw, (70, 610, 390, 850), "EDF7F1", "25835B", "Joint-space samples", ["6-DoF joint limits", "strict self-collision", "normalized q"])
    rounded_box(draw, (510, 570, 920, 890), "EDF7F1", "25835B", "Reachability graph G_R", ["GNG, guarded GNG,", "or Halton/PRM", "node = EEF pose + full q", "shared validated k-NN"])
    rounded_box(draw, (1050, 270, 1375, 505), "FFF7E6", "A66B00", "Graph intersection", ["target node within", "50 mm of reachable", "start component"])
    rounded_box(draw, (1050, 625, 1375, 870), "FFF7E6", "A66B00", "Fast graph query", ["capsule invalidation", "Dijkstra shortest path", "block obstacle edges"])
    rounded_box(draw, (1490, 380, 1740, 700), "F8ECF3", "8F3B69", "Exact validator", ["MoveIt/FCL meshes", "interpolated states", "reject edge + replan", "preview path only"])

    arrow(draw, (390, 300), (510, 300), BLUE)
    arrow(draw, (390, 730), (510, 730), "25835B")
    arrow(draw, (920, 300), (1050, 360), BLUE)
    arrow(draw, (920, 730), (1050, 750), "25835B")
    arrow(draw, (1210, 505), (1210, 625), "A66B00")
    arrow(draw, (1375, 750), (1490, 610), "8F3B69")
    arrow(draw, (1490, 460), (1375, 430), "8F3B69", 5)
    draw.text((1390, 415), "failed edge", font=font(19, True), fill=pc("8F3B69"))
    draw.rounded_rectangle((1450, 820, 1770, 945), radius=18, fill=pc("FFF0F0"), outline=pc("9B1C1C"), width=3)
    draw.text((1480, 842), "Safety boundary", font=font(25, True), fill=pc("9B1C1C"))
    draw.text((1480, 880), "No controller publisher", font=font(20), fill=pc(BLACK))
    draw.text((1480, 908), "No action client", font=font(20), fill=pc(BLACK))
    image.save(path, dpi=(180, 180))


def mean_sd(rows, method, field, predicate=None):
    values = [
        float(r[field])
        for r in rows
        if r["method"] == method and (predicate is None or predicate(r))
    ]
    return statistics.mean(values), statistics.stdev(values)


def bar(draw, x, y_bottom, width, height, color, label, value, max_value):
    top = y_bottom - int(height * value / max_value)
    draw.rounded_rectangle((x, top, x + width, y_bottom), radius=8, fill=pc(color))
    draw.text((x + width / 2, top - 35), f"{value:.1f}", anchor="mm", font=font(20, True), fill=pc(INK))
    draw.text((x + width / 2, y_bottom + 28), label, anchor="mm", font=font(18), fill=pc(BLACK))


def panel(draw, xy, title, entries, max_value, unit):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=pc("FAFBFC"), outline=pc("D3D9E0"), width=3)
    draw.text((x1 + 25, y1 + 20), title, font=font(25, True), fill=pc(INK))
    baseline = y2 - 78
    draw.line((x1 + 38, baseline, x2 - 30, baseline), fill=pc("9AA5B1"), width=2)
    n = len(entries)
    available = x2 - x1 - 100
    bw = min(105, int(available / (n * 1.8)))
    gap = available / n
    for i, (label, value, color) in enumerate(entries):
        x = int(x1 + 55 + i * gap + (gap - bw) / 2)
        bar(draw, x, baseline, bw, y2 - y1 - 175, color, label, value, max_value)
    draw.text((x1 + 22, y2 - 35), unit, font=font(17), fill=pc(MUTED))


def build_results_figure(path: Path, rows):
    image = Image.new("RGB", (1800, 1120), pc(WHITE))
    draw = ImageDraw.Draw(image)
    draw.text((70, 32), "Development selection and held-out fixed-task benchmark", font=font(36, True), fill=pc(INK))
    draw.text((70, 82), "Held-out variability is over 50 paired deterministic sample-stream offsets per method.", font=font(22), fill=pc(MUTED))
    methods = (("GNG", "gng", "2E74B5"), ("Guarded", "guarded_gng", "25835B"), ("Halton", "halton_prm", "AAB7C4"))
    success = []
    builds = []
    components = []
    for label, method, color in methods:
        method_rows = [r for r in rows if r["method"] == method]
        success.append((label, 100.0 * sum(r["dynamic_valid"] == "True" for r in method_rows) / len(method_rows), color))
        builds.append((label, statistics.median(float(r["build_time_ms"]) for r in method_rows), color))
        components.append((label, statistics.median(float(r["components"]) for r in method_rows), color))
    development = [("0%", 0.0, "738394"), ("10%", 12.5, "5D879F"), ("25%", 50.0, "3E7E7B"), ("50%", 87.5, "2F765B"), ("75%", 100.0, "25835B")]
    panel(draw, (45, 145, 880, 570), "Development failure-offset recovery", development, 100, "percent of eight selected failures; selection only")
    panel(draw, (920, 145, 1755, 570), "Held-out obstacle-update success", success, 100, "percent; Wilson intervals in Table III")
    panel(draw, (45, 625, 880, 1050), "Held-out connected components", components, 15, "median; lower component count only")
    panel(draw, (920, 625, 1755, 1050), "Held-out roadmap build", builds, 1900, "median milliseconds; lower is better")
    image.save(path, dpi=(180, 180))


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9.5, italic=True, color=MUTED)
    return p


def add_equation(doc, text, number):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{text}                                      ({number})")
    set_run_font(r, name="Cambria Math", size=10.5)
    return p


def add_reference(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(f"[{number}] {text}")
    set_run_font(r, size=9.2)


def add_callout(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "7")
        border.set(qn("w:color"), "AAB7C4")
        borders.append(border)
    p_pr.append(borders)
    r = p.add_run(f"{title}: ")
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r)


def style_document(doc):
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Working draft | Page ")
    set_run_font(r, size=8.5, color=MUTED)
    add_page_number(p)


def add_table_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=9.3, color=BLACK):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold, color=color)


def fmt(mean, sd, digits=2):
    return f"{mean:.{digits}f} +/- {sd:.{digits}f}"


def build_legacy_document():
    ASSETS.mkdir(exist_ok=True)
    QA.mkdir(exist_ok=True)
    rows = list(csv.DictReader((SOURCE / "pilot.csv").open(newline="", encoding="utf-8")))
    architecture = ASSETS / "system_architecture.png"
    results = ASSETS / "pilot_results.png"
    build_architecture_figure(architecture)
    build_results_figure(results, rows)

    doc = Document()
    style_document(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("Topology-Preserving End-Effector Reachability Roadmaps with Two-Stage Collision Validation")
    set_run_font(r, size=21, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("[AUTHOR NAME], [CO-AUTHOR NAME], and [SUPERVISOR NAME]")
    set_run_font(r, size=11.5, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("[Laboratory / Department / University], [City, Country] | [corresponding@email]")
    set_run_font(r, size=10, italic=True, color=MUTED)

    add_callout(
        doc,
        "Draft status",
        "Working manuscript based on the implementation and pilot data frozen on 24 August 2026. "
        "The current evaluation uses three deterministic seeds per method and is not yet an ICRA-ready statistical study.",
    )

    h = doc.add_heading("Abstract", level=1)
    h.paragraph_format.space_before = Pt(10)
    abstract = (
        "We present a dual-topology planning architecture for target-directed manipulation in which a learned environment graph is intersected with a configuration-aware end-effector reachability graph. "
        "The environment is represented by a Dynamic Density Growing Neural Gas (DD-GNG) graph built from wrist-mounted RGB-D observations; semantically labelled nodes act as target hypotheses, while unlabelled nodes and edges define obstacle geometry. "
        "The robot roadmap is learned in normalized six-dimensional joint space using Growing Neural Gas (GNG), while retaining the full joint configuration and end-effector pose at every node. "
        "Runtime updates first invalidate roadmap nodes and edges using cached swept-body capsules, then validate the shortlisted path against exact robot meshes and DD-GNG-derived geometry in a MoveIt/FCL PlanningScene. "
        "A rejected exact-collision edge is disabled and graph search is repeated. In a preliminary controller-free benchmark with 800 nodes and three seeds per method, GNG produced 2.0 +/- 1.0 connected components versus 12.0 +/- 3.46 for a matched Halton/PRM baseline. "
        "Mean clear-scene query time was 6.89 ms versus 13.14 ms, and dynamic-obstacle query time was 13.86 ms versus 35.19 ms. All twelve method-seed-scenario queries returned exact-validated preview paths. "
        "These results are promising but descriptive: the current pilot does not establish statistical significance, test the exact-rejection replan branch, or evaluate physical trajectory execution."
    )
    doc.add_paragraph(abstract)
    p = doc.add_paragraph()
    r = p.add_run("Keywords - ")
    set_run_font(r, bold=True, italic=True)
    r = p.add_run("reachability graph, Growing Neural Gas, manipulation planning, dynamic obstacles, collision checking, topological planning")
    set_run_font(r, italic=True)

    doc.add_heading("I. Introduction", level=1)
    doc.add_paragraph(
        "A manipulation system operating from live RGB-D observations must answer two coupled questions: where can the robot end effector reach, and which of those reachable configurations remain connected to the current arm state after the environment changes? Conventional single-query planners answer these questions jointly for each request, whereas roadmap methods amortize configuration-space exploration across multiple queries [2]. A persistent roadmap is attractive for interactive manipulation, but a Cartesian point alone is not an executable robot state: multiple inverse-kinematic branches may place the end effector at nearly the same position while belonging to different collision-free components."
    )
    doc.add_paragraph(
        "This work investigates a dual-topology alternative. A perception topology G_E compresses the observed workspace into semantically labelled target nodes and obstacle-bearing nodes and edges. A robot topology G_R compresses valid joint configurations while retaining the complete configuration associated with each end-effector pose. A target becomes actionable only when its environment node lies within a tolerance of a reachable node that belongs to the component attached to the measured robot state. The output is deliberately a preview joint path rather than an execution command."
    )
    doc.add_paragraph("The current implementation makes three contributions:")
    for item in (
        "A configuration-retaining GNG reachability roadmap that learns the topology of valid normalized joint samples and preserves multiple inverse-kinematic branches at similar Cartesian locations.",
        "A fast environment-update mechanism that compares DD-GNG obstacle nodes and edges with cached swept capsules for the complete moving body, avoiding roadmap reconstruction after every perception update.",
        "A two-stage safety gate in which capsule-safe paths are interpolated and checked with MoveIt/FCL meshes; rejected roadmap edges are persisted for the current environment and Dijkstra search is repeated.",
    ):
        doc.add_paragraph(item, style="List Number")
    doc.add_paragraph(
        "The preliminary evaluation compares GNG against an equal-budget deterministic Halton/PRM construction. The manuscript intentionally separates observed pilot evidence from planned evaluation."
    )

    doc.add_heading("II. Related Work", level=1)
    doc.add_heading("A. Sampling-Based Motion Planning", level=2)
    doc.add_paragraph(
        "Probabilistic roadmaps (PRMs) preprocess collision-free configurations and feasible local connections, then answer subsequent queries through graph search [2]. Rapidly exploring random trees (RRTs) provide an influential single-query alternative with strong exploratory bias in high-dimensional spaces [3]. Our current baseline is a deterministic Halton/PRM variant because it shares the same node budget, k-nearest connection rule, edge interpolation, and collision predicate as the proposed GNG roadmap. RRTConnect remains a required baseline for the full study."
    )
    doc.add_heading("B. Growing Neural Gas and Topological Roadmaps", level=2)
    doc.add_paragraph(
        "GNG incrementally introduces units and connections to learn topological relations in an input distribution [1]. Prior work has used GNG for topological environment maps [7] and topology-informed navigation roadmaps over uncertain occupancy maps [6]. In contrast, the present robot graph learns a distribution of valid manipulator joint configurations, stores the corresponding end-effector pose, and is intersected online with a separately learned RGB-D environment topology. This distinction is important: environment-space connectivity alone does not encode manipulator joint limits, self-collision, or disconnected inverse-kinematic branches."
    )
    doc.add_heading("C. Collision Checking in Manipulation", level=2)
    doc.add_paragraph(
        "MoveIt provides a reusable planning architecture for complex robot systems [4], while the Flexible Collision Library (FCL) supplies collision and proximity queries over articulated and mesh geometry [5]. Exact geometry is valuable but expensive when applied indiscriminately to every roadmap edge after each environment update. Our design therefore uses conservative cached capsule sweeps as a broad phase and reserves mesh-level FCL validation for the shortlisted path."
    )

    doc.add_heading("III. Problem Formulation", level=1)
    doc.add_paragraph(
        "Let Q be the six-dimensional joint-limit box and Q_free be the subset satisfying joint limits and strict self-collision constraints. Each robot-graph node v_i stores a normalized configuration q_i in Q_free, the physical joint vector, and a forward-kinematic end-effector pose x_i = FK(q_i). The normalization uses each joint range so that no joint dominates nearest-neighbour selection solely because of units or span."
    )
    add_equation(doc, "d_Q(q_i,q_j) = sqrt( sum_k ((q_i,k - q_j,k)/(q_k^max-q_k^min))^2 )", 1)
    doc.add_paragraph(
        "The environment graph G_E = (V_E, E_E) is expressed in the world frame. A semantic label partitions V_E into target candidates T and obstacle samples O; unlabelled environment edges are treated as obstacle segments. For target t in T, the intersection candidate set contains robot nodes within radius r_I that remain unblocked and connected to the node attached to the measured configuration q_0."
    )
    add_equation(doc, "I(t) = { v_i in V_R | ||p_i-p_t||_2 <= r_I, free(v_i), connected(v_i,v_start) }", 2)
    doc.add_paragraph(
        "The planner selects a candidate in I(t) and minimizes accumulated roadmap edge cost using Dijkstra search. A plan is considered valid only when every interpolated state on the resulting route passes the exact second-stage collision test. The current system does not claim dynamic or kinodynamic feasibility because velocity, acceleration, torque, and closed-loop execution are outside the present preview interface."
    )

    doc.add_heading("IV. Method", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(architecture), width=Inches(6.35))
    doc_pr = run._r.xpath(".//wp:docPr")
    if doc_pr:
        doc_pr[0].set("descr", "Architecture diagram of the dual-topology reachability pipeline")
    add_caption(doc, "Fig. 1. Dual-topology pipeline. Environment and reachability graphs meet at a target intersection; capsule filtering precedes exact MoveIt/FCL validation. The output remains preview-only.")

    doc.add_heading("A. Environment Topology", level=2)
    doc.add_paragraph(
        "A wrist-mounted Intel RealSense D405 provides aligned depth and colour. Depth samples are deprojected using live intrinsics and transformed into the world frame using tf2 at the frame timestamp. DD-GNG maintains a compact graph whose stable node identifiers, three-dimensional positions, semantic class, confidence, and edges are published as a typed ROS 2 message. Robot-body points are removed before graph learning using the same link capsule model used by the broad-phase planner. Semantically labelled nodes form target hypotheses; all other nodes and edges are collision geometry, except for a small exclusion region around the selected target."
    )

    doc.add_heading("B. Configuration-Aware GNG Roadmap", level=2)
    doc.add_paragraph(
        "Candidate joint vectors are generated inside URDF limits and rejected if MoveIt forward kinematics fails or strict self-collision is detected. GNG is trained over normalized valid configurations and grows to the requested node budget. Each learned prototype is mapped back to physical joints and stores FK pose plus the complete six-joint vector. For the controlled baseline, a Halton sequence generates the same number of accepted nodes. Both methods then share an identical connection stage: ten nearest neighbours are considered, normalized joint distance is capped at 0.75, Cartesian end-effector separation is capped at 0.14 m, and every candidate edge is interpolated and collision checked."
    )

    doc.add_heading("C. Cached Full-Body Invalidation", level=2)
    doc.add_paragraph(
        "For every roadmap node and edge, the system caches capsules spanning link2 through link7, the end effector, both fingers, and the D405 payload over interpolated configurations. The fixed-base link1 segment is omitted by default so that a known support surface does not invalidate the whole roadmap. After an environment update, point-to-segment and segment-to-segment distance tests mark blocked nodes and edges without repeating forward kinematics or reconstructing G_R."
    )
    add_equation(doc, "blocked(c,o) iff distance(axis(c), o) < radius(c) + clearance", 3)

    doc.add_heading("D. Exact Mesh Validation and Replanning", level=2)
    doc.add_paragraph(
        "The broad phase produces a candidate route. The environment graph is then materialized in a MoveIt PlanningScene: nodes become spheres and edges become cylinders aligned with their three-dimensional segments. The robot collision model includes the fused wrist bracket and D405 mesh. Every route edge is interpolated at a normalized-joint step of 0.05 and checked against the exact robot and environment geometry using FCL. On collision, the offending roadmap edge is disabled, Dijkstra search is repeated, and the process continues for at most twenty replans. Exact-blocked edges persist until the environment graph changes."
    )
    doc.add_paragraph(
        "A valid plan message reports graph method, start and goal identifiers, blocked-node and blocked-edge counts, graph cost, planning time, exact state checks, exact replan count, exact-validation time, a joint-path preview, and an end-effector path. It does not publish to a FollowJointTrajectory controller or create a controller action client."
    )

    doc.add_heading("V. Implementation and Parameters", level=1)
    doc.add_paragraph(
        "The implementation runs as ROS 2 Humble C++ nodes on the robot's AGX computer. MoveIt RobotModel/RobotState provide kinematics and self-collision checks; MoveIt PlanningScene and FCL provide exact environment collision checking. Perception and reachability are separate processes so the roadmap benchmark can run without opening the camera, MoveGroup, controller manager, or hardware interface. The live deployment at the data freeze used GNG with 800 nodes, 4518 edges, three connected components, and a measured build time of approximately 1.72 s."
    )
    table_i_caption = add_caption(doc, "Table I. Principal parameters used for the current pilot.")
    keep_with_next(table_i_caption)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [3000, 1860, 4500]
    set_table_geometry(table, widths)
    headers = ("Parameter", "Value", "Role")
    for idx, value in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], LIGHT_BLUE)
        add_table_text(table.rows[0].cells[idx], value, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    parameter_rows = [
        ("Roadmap nodes", "800", "Equal node budget for GNG and Halton/PRM"),
        ("GNG training samples", "4000", "Valid normalized joint-space signals"),
        ("GNG epochs / insertion", "4 / 20", "Growth schedule"),
        ("Neighbour candidates", "10", "Shared k-NN connection stage"),
        ("Max normalized joint distance", "0.75", "Reject long local connections"),
        ("Max Cartesian edge length", "0.14 m", "Restrict EEF displacement"),
        ("Target intersection radius", "0.05 m", "Environment-to-robot graph overlap"),
        ("Obstacle clearance", "0.035 m", "Capsule broad-phase margin"),
        ("Body interpolation step", "0.08", "Cached capsule edge sweep"),
        ("Exact interpolation step", "0.05", "MoveIt/FCL route validation"),
        ("Environment primitive radii", "0.012 / 0.006 m", "Point spheres / edge cylinders"),
        ("Maximum exact replans", "20", "Bound edge rejection and graph re-search"),
    ]
    for row_values in parameter_rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            add_table_text(cells[idx], value, align=WD_ALIGN_PARAGRAPH.CENTER if idx == 1 else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(table, widths)
    p = doc.add_paragraph("Source: current topo_gng.yaml configuration snapshot.")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.runs[0], size=8.5, italic=True, color=MUTED)

    doc.add_heading("VI. Preliminary Evaluation", level=1)
    doc.add_heading("A. Questions and Protocol", level=2)
    doc.add_paragraph(
        "The pilot asks whether a GNG-quantized roadmap improves connectivity and repeated query latency relative to an equal-size deterministic Halton/PRM baseline. For each method and seed, an isolated ROS domain launches only the reachability node, publishes a synthetic home joint state, selects a target at a known roadmap node, and records a clear-scene plan. It then inserts a synthetic obstacle on the original route and records the updated plan. The experiment uses seeds 0-2, yielding six roadmap builds and twelve planning queries. This design verifies repeatability and message instrumentation but is too small for inferential statistics."
    )

    doc.add_heading("B. Pilot Results", level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(results), width=Inches(6.35))
    doc_pr = run._r.xpath(".//wp:docPr")
    if doc_pr:
        doc_pr[0].set("descr", "Bar charts comparing preliminary GNG and Halton PRM benchmark results")
    add_caption(doc, "Fig. 2. Descriptive pilot means. GNG required more roadmap construction time but produced fewer connected components and lower query latency.")

    table_ii_caption = add_caption(doc, "Table II. Pilot results, mean +/- sample standard deviation over three seeds.")
    keep_with_next(table_ii_caption)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [4400, 2480, 2480]
    set_table_geometry(table, widths)
    for idx, value in enumerate(("Metric", "GNG", "Halton/PRM")):
        shade_cell(table.rows[0].cells[idx], LIGHT_BLUE)
        add_table_text(table.rows[0].cells[idx], value, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    result_rows = [
        ("Roadmap nodes", "800", "800"),
        ("Roadmap edges", fmt(*mean_sd(rows, "gng", "edges"), 1), fmt(*mean_sd(rows, "halton_prm", "edges"), 1)),
        ("Connected components", fmt(*mean_sd(rows, "gng", "components"), 2), fmt(*mean_sd(rows, "halton_prm", "components"), 2)),
        ("Build time (ms)", fmt(*mean_sd(rows, "gng", "build_time_ms"), 2), fmt(*mean_sd(rows, "halton_prm", "build_time_ms"), 2)),
        ("Clear planning time (ms)", fmt(*mean_sd(rows, "gng", "clear_planning_time_ms"), 2), fmt(*mean_sd(rows, "halton_prm", "clear_planning_time_ms"), 2)),
        ("Clear exact state checks", fmt(*mean_sd(rows, "gng", "clear_exact_checks"), 1), fmt(*mean_sd(rows, "halton_prm", "clear_exact_checks"), 1)),
        ("Clear exact time (ms)", fmt(*mean_sd(rows, "gng", "clear_exact_time_ms"), 2), fmt(*mean_sd(rows, "halton_prm", "clear_exact_time_ms"), 2)),
        ("Dynamic blocked nodes", fmt(*mean_sd(rows, "gng", "dynamic_blocked_nodes"), 1), fmt(*mean_sd(rows, "halton_prm", "dynamic_blocked_nodes"), 1)),
        ("Dynamic blocked edges", fmt(*mean_sd(rows, "gng", "dynamic_blocked_edges"), 1), fmt(*mean_sd(rows, "halton_prm", "dynamic_blocked_edges"), 1)),
        ("Dynamic planning time (ms)", fmt(*mean_sd(rows, "gng", "dynamic_planning_time_ms"), 2), fmt(*mean_sd(rows, "halton_prm", "dynamic_planning_time_ms"), 2)),
        ("Dynamic exact state checks", fmt(*mean_sd(rows, "gng", "dynamic_exact_checks"), 1), fmt(*mean_sd(rows, "halton_prm", "dynamic_exact_checks"), 1)),
        ("Dynamic exact time (ms)", fmt(*mean_sd(rows, "gng", "dynamic_exact_time_ms"), 2), fmt(*mean_sd(rows, "halton_prm", "dynamic_exact_time_ms"), 2)),
        ("Clear / dynamic success", "3/3 / 3/3", "3/3 / 3/3"),
        ("Observed exact replans", "0", "0"),
    ]
    for label, gng, halton in result_rows:
        cells = table.add_row().cells
        add_table_text(cells[0], label)
        add_table_text(cells[1], gng, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_table_text(cells[2], halton, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(table, widths)
    p = doc.add_paragraph("Source: reachability_benchmark_pilot_20260824.csv. These are descriptive pilot statistics; n = 3 roadmaps per method.")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.runs[0], size=8.5, italic=True, color=MUTED)

    doc.add_paragraph(
        "Under the equal node budget, GNG generated approximately 39% more validated edges and reduced the mean component count from 12 to 2. Its roadmap build was approximately 21% slower, consistent with an additional learning stage. In return, mean clear-scene query time was approximately 48% lower and mean dynamic-obstacle query time approximately 61% lower than Halton/PRM. GNG also required fewer exact state checks in both conditions. These comparisons are effect-size descriptions of the six-run pilot, not evidence of population-level superiority."
    )
    doc.add_paragraph(
        "All twelve queries produced paths with exact_collision_valid=true. However, exact_replans remained zero: the injected obstacle was already conservative enough to be captured by the capsule broad phase. Consequently, the pilot confirms that exact validation runs on candidate trajectories but does not experimentally exercise the branch that rejects a capsule-safe edge and reruns Dijkstra."
    )

    doc.add_heading("VII. Limitations and Planned Evaluation", level=1)
    limitations = [
        ("Statistical scale", "Increase to at least 50 deterministic seeds per method and report confidence intervals, paired effect sizes, and nonparametric significance tests where justified."),
        ("Planner baselines", "Add MoveIt/OMPL RRTConnect and, if computationally feasible, PRM* using the same collision model, start-goal pairs, and wall-clock limits."),
        ("Exact-replan stress case", "Generate thin or offset obstacles that pass the capsule broad phase but intersect a robot or D405 mesh, then measure edge rejection count and recovery latency."),
        ("Ablation", "Compare capsule-only, FCL-only, and two-stage validation; vary 200, 400, 800, and 1600 roadmap nodes; remove environment edges; and disable strict collision handling."),
        ("Physical validation", "Measure predicted-versus-observed reachability and collision-free execution across repeated real-robot trials with velocity, acceleration, emergency-stop, and human safety procedures."),
        ("Dynamic validity", "The current graph invalidates updated geometry but does not predict obstacle motion or enforce time-parameterized clearance. Dynamic-scene claims must remain limited to re-querying after graph updates."),
        ("Perception uncertainty", "DD-GNG geometry is treated deterministically after construction. Future work should propagate depth, segmentation, and transform uncertainty into collision margins or probabilistic risk."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    widths = [2200, 7160]
    set_table_geometry(table, widths)
    for idx, value in enumerate(("Gap", "Required experiment or change")):
        shade_cell(table.rows[0].cells[idx], LIGHT_GRAY)
        add_table_text(table.rows[0].cells[idx], value, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for label, detail in limitations:
        cells = table.add_row().cells
        add_table_text(cells[0], label, bold=True)
        add_table_text(cells[1], detail)
    set_table_geometry(table, widths)

    doc.add_heading("VIII. Conclusion", level=1)
    doc.add_paragraph(
        "This draft introduced a dual-topology manipulation architecture in which a learned semantic environment graph intersects a configuration-aware GNG reachability roadmap. Cached full-body capsule sweeps support rapid invalidation after environment updates, while a MoveIt/FCL second stage validates exact robot geometry and can remove offending roadmap edges before replanning. A small deterministic pilot suggests that the GNG roadmap trades higher construction time for substantially improved connectivity and lower repeated-query latency relative to an equal-size Halton/PRM baseline. The result is technically encouraging but not publication-conclusive. The immediate path to a defensible ICRA submission is a preregistered large-seed benchmark, RRTConnect comparison, exact-rejection stress test, ablation study, and controlled real-robot validation."
    )

    doc.add_heading("References", level=1)
    references = [
        "B. Fritzke, 'A Growing Neural Gas Network Learns Topologies,' in Advances in Neural Information Processing Systems 7, pp. 625-632, 1995.",
        "L. E. Kavraki, P. Svestka, J.-C. Latombe, and M. H. Overmars, 'Probabilistic Roadmaps for Path Planning in High-Dimensional Configuration Spaces,' IEEE Transactions on Robotics and Automation, vol. 12, no. 4, pp. 566-580, 1996, doi: 10.1109/70.508439.",
        "S. M. LaValle, 'Rapidly-Exploring Random Trees: A New Tool for Path Planning,' Technical Report 98-11, Iowa State University, 1998.",
        "D. Coleman, I. A. Sucan, S. Chitta, and N. Correll, 'Reducing the Barrier to Entry of Complex Robotic Software: a MoveIt! Case Study,' Journal of Software Engineering for Robotics, vol. 5, no. 1, pp. 3-16, 2014, doi: 10.6092/JOSER_2014_05_01_p3.",
        "J. Pan, S. Chitta, and D. Manocha, 'FCL: A General Purpose Library for Collision and Proximity Queries,' in Proc. IEEE International Conference on Robotics and Automation, 2012, doi: 10.1109/ICRA.2012.6225337.",
        "M. Saroya, G. Best, and G. A. Hollinger, 'Roadmap Learning for Probabilistic Occupancy Maps With Topology-Informed Growing Neural Gas,' IEEE Robotics and Automation Letters, 2021, doi: 10.1109/LRA.2021.3068886.",
        "Y. Toda et al., 'Growing Neural Gas Based Topological Environmental Map Building and Path Planning in Unknown Environment,' Journal of Japan Society for Fuzzy Theory and Intelligent Informatics, vol. 33, no. 4, pp. 872-884, 2021, doi: 10.3156/jsoft.33.4_872.",
    ]
    for idx, reference in enumerate(references, 1):
        add_reference(doc, idx, reference)

    doc.core_properties.title = "Topology-Preserving End-Effector Reachability Roadmaps with Two-Stage Collision Validation"
    doc.core_properties.subject = "Preliminary ICRA-oriented working manuscript"
    doc.core_properties.keywords = "GNG, reachability, manipulation, collision checking, ROS 2"
    doc.core_properties.comments = "Draft generated from a three-seed pilot; not for submission without expanded evaluation."
    doc.save(OUTPUT)
    print(OUTPUT)


def build_document():
    ASSETS.mkdir(exist_ok=True)
    QA.mkdir(exist_ok=True)
    benchmark_path = DATA / "reachability_benchmark_heldout_50_three_methods_20260823.csv"
    rows = list(csv.DictReader(benchmark_path.open(newline="", encoding="utf-8")))
    architecture = ASSETS / "system_architecture_heldout.png"
    results = ASSETS / "heldout_results.png"
    build_architecture_figure(architecture)
    build_results_figure(results, rows)

    doc = Document()
    style_document(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("Target-Robustness–Connectivity Trade-offs in GNG-Quantized Reachability Roadmaps with Two-Stage Collision Validation")
    set_run_font(r, size=20.5, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("[AUTHOR NAME], [CO-AUTHOR NAME], and [SUPERVISOR NAME]")
    set_run_font(r, size=11.5, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("[Laboratory / Department / University], [City, Country] | [corresponding@email]")
    set_run_font(r, size=10, italic=True, color=MUTED)

    add_callout(
        doc,
        "Draft status",
        "Working manuscript frozen on 23 August 2026. The reported experiment comprises 150 independently built roadmaps and 300 controller-free planning queries. It is a held-out fixed-task study, not yet a multi-scene or physical-robot evaluation.",
    )

    h = doc.add_heading("Abstract", level=1)
    h.paragraph_format.space_before = Pt(10)
    doc.add_paragraph(
        "We study persistent end-effector reachability roadmaps for manipulation under obstacle-graph updates. A dual-topology ROS 2 architecture intersects semantic target nodes from a perception graph with a robot graph whose nodes retain complete six-joint configurations. Candidate routes are filtered by cached full-body capsule sweeps and then validated with interpolated MoveIt/FCL mesh checks; rejected edges are disabled and graph search is repeated. We compare three equal-size 800-node constructions: a Growing Neural Gas (GNG)-quantized k-nearest-neighbour roadmap, a guarded variant that reserves 75% of its node budget for deterministic low-discrepancy samples, and a Halton/PRM baseline. The guard fraction was chosen on eight development failures and frozen before evaluation on 50 paired held-out sample-stream offsets. On one fixed anchored start-target query with a synthetic midpoint point obstacle, obstacle-update success was 46/50 for GNG (92%, Wilson 95% CI 81.2–96.8), 49/50 for guarded GNG (98%, 89.5–99.6), and 49/50 for Halton/PRM (98%, 89.5–99.6); the global difference was not significant (Cochran Q(2)=3.00, p=.223). Guarded GNG improved the paired success estimate over GNG by 6 percentage points (paired-bootstrap 95% CI -2 to 14; Holm-adjusted exact McNemar p=.750) and tied Halton/PRM (0 points, -6 to 6; p=1.000). Pure GNG produced a median of one connected component versus twelve for both alternatives, yet failed more often after local invalidation. The fixed-task results suggest a target-robustness–connectivity trade-off rather than method superiority or general reachable-area coverage."
    )
    p = doc.add_paragraph()
    r = p.add_run("Keywords - ")
    set_run_font(r, bold=True, italic=True)
    r = p.add_run("reachability roadmap, Growing Neural Gas, low-discrepancy sampling, manipulation planning, collision checking, ROS 2")
    set_run_font(r, italic=True)

    doc.add_heading("I. Introduction", level=1)
    doc.add_paragraph(
        "A manipulation system that reasons from RGB-D observations must distinguish geometric reachability from executable connectivity. A Cartesian point may be reachable through several inverse-kinematic branches, yet only some branches can be joined collision-free to the measured arm configuration. Capability and reachability maps summarize workspace access [4], [5], while probabilistic roadmaps amortize configuration-space exploration across queries [2]. Neither representation alone directly answers whether a semantically selected scene target remains connected after the observed obstacle graph changes."
    )
    doc.add_paragraph(
        "We investigate a dual-topology architecture. An environment topology G_E compresses the observed workspace into target-labelled nodes and obstacle-bearing nodes and edges. A robot topology G_R stores collision-valid joint configurations, their end-effector poses, and validated local transitions. A target becomes actionable only when an overlapping G_R node lies in the component connected to the measured start. The implementation is deliberately preview-only: it publishes a joint trajectory on a private visualization topic and has no controller publisher or FollowJointTrajectory action client."
    )
    doc.add_paragraph(
        "The initial hypothesis was that GNG's distribution-adaptive prototypes would improve roadmap topology. Source inspection revealed an important qualification: the current implementation retains GNG prototypes but discards the learned GNG adjacency before constructing a shared k-nearest-neighbour graph. We therefore use the precise term GNG-quantized k-NN roadmap and do not claim that Fritzke's learned edges are preserved [1]. The benchmark further exposed a tension between global graph connectivity and local support near an obstacle-conditioned target."
    )
    doc.add_paragraph("This working paper contributes:")
    for item in (
        "A typed dual-topology planning architecture that retains full configurations at robot nodes and intersects them with semantic environment targets.",
        "A two-stage obstacle-update mechanism combining cached swept-body capsules with lazy interpolated MoveIt/FCL validation and persistent edge rejection.",
        "A guarded-GNG construction that mixes learned prototypes with deterministic raw valid samples under a fixed node budget.",
        "A development/held-out-offset evaluation over 150 roadmap builds with paired confidence intervals and nonparametric tests, revealing a target-robustness–connectivity hypothesis rather than an unsupported superiority claim.",
    ):
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("II. Related Work", level=1)
    doc.add_heading("A. Roadmaps, Low-Discrepancy Sampling, and GNG", level=2)
    doc.add_paragraph(
        "PRM separates an offline roadmap phase from online start-goal queries [2]. Low-discrepancy sequences, including Halton sampling, provide deterministic space-filling alternatives to pseudorandom sampling [3]. GNG incrementally learns prototypes and adjacency for an input distribution [1]. Our comparison isolates the effect of node generation: all three methods share joint limits, validity predicates, an 800-node cap, k-nearest-neighbour construction, edge-length thresholds, and interpolated edge checking. Because the GNG adjacency is rebuilt, this study evaluates adaptive quantization rather than the complete original GNG graph."
    )
    doc.add_heading("B. Reachability and Experience Representations", level=2)
    doc.add_paragraph(
        "Capability maps represent directional end-effector access throughout the workspace [4], and reachability maps support manipulation-task analysis [5]. Sparse roadmap spanners preserve coverage and path quality while limiting graph growth [9]. E-Graphs and Thunder reuse previous motion experience [10], [11]. Our robot graph differs from a binary reachability volume by preserving configuration-space transitions and multiple branches at similar end-effector positions; however, it currently provides no asymptotic optimality or completeness guarantee."
    )
    doc.add_heading("C. Environment Updates and Learned Bias", level=2)
    doc.add_paragraph(
        "Lazy PRM postpones expensive validation until a candidate path is queried [6]. Dynamic Roadmaps cache how workspace regions affect roadmap elements [7], while changing-environment PRM variants invalidate and reinforce affected regions [8]. Our capsule cache follows this broad design pattern and uses exact FCL checks only on shortlisted routes. Learned sampling can accelerate difficult planning problems, but mixtures with globally supportive sampling reduce the risk of an overly concentrated learned bias [12]–[14]. Guarded GNG is a deterministic instance of that principle, not a claim of probabilistic completeness. MoveIt supplies robot modelling and planning-scene infrastructure [15], and FCL supplies exact collision/proximity queries [16]."
    )

    doc.add_heading("III. Problem Formulation", level=1)
    doc.add_paragraph(
        "Let Q be the six-dimensional joint-limit box and Q_free the subset accepted by the robot validity predicate. A robot node v_i stores the physical joint vector q_i, its range-normalized form, and end-effector pose x_i=FK(q_i). Normalization prevents a joint with a wider numerical range from dominating neighbour selection."
    )
    add_equation(doc, "d_Q(q_i,q_j) = sqrt( sum_k ((q_i,k-q_j,k)/(q_k^max-q_k^min))^2 )", 1)
    doc.add_paragraph(
        "The environment graph G_E=(V_E,E_E) is expressed in the world frame. Semantic labels partition target candidates T from obstacle primitives O. For target t, the query considers unblocked robot nodes within Cartesian intersection radius r_I that lie in the start component."
    )
    add_equation(doc, "I(t) = {v_i in V_R | ||p_i-p_t||_2 <= r_I, free(v_i), connected(v_i,v_start)}", 2)
    doc.add_paragraph(
        "Dijkstra search minimizes accumulated roadmap cost to a candidate in I(t). A returned preview is valid only if every interpolated route state passes exact self- and environment-collision checks. The system does not claim kinodynamic feasibility: time parameterization, torque limits, closed-loop tracking, moving-obstacle prediction, and execution are outside the present interface."
    )

    doc.add_heading("IV. System and Roadmap Construction", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(architecture), width=Inches(6.35))
    doc_pr = run._r.xpath(".//wp:docPr")
    if doc_pr:
        doc_pr[0].set("descr", "Dual-topology reachability pipeline with guarded GNG and two-stage collision validation")
    add_caption(doc, "Fig. 1. Dual-topology pipeline. Environment targets intersect configuration-retaining robot nodes; capsule invalidation precedes exact MoveIt/FCL route validation. Output is preview-only.")

    doc.add_heading("A. Environment and Robot Topologies", level=2)
    doc.add_paragraph(
        "In the full system, a wrist-mounted RealSense D405 is deprojected with live depth intrinsics and transformed into the world frame at the image timestamp. DD-GNG maintains stable three-dimensional nodes, semantic class, confidence, and edges. Target-labelled nodes are goal hypotheses; other nodes and segments become collision geometry, except for a small exclusion region around the selected target. Perception and reachability run as separate ROS 2 processes so roadmap experiments can be isolated from the camera and hardware interfaces. The present benchmark injects a synthetic target and obstacle and therefore does not evaluate DD-GNG perception."
    )
    doc.add_heading("B. Three Equal-Budget Roadmaps", level=2)
    doc.add_paragraph(
        "Valid configurations are first generated within URDF limits and rejected on failed forward kinematics or strict self-collision. Pure GNG is trained on 4000 valid normalized signals, grows prototypes toward the node cap, and maps prototypes back to physical joint configurations. Halton/PRM accepts deterministic low-discrepancy valid configurations directly. Both are subsequently connected by the same k-NN stage."
    )
    doc.add_paragraph(
        "Guarded GNG reserves 75% of the same 800-node budget for deterministic raw valid samples selected by stratified indices across the valid-sample stream; the remainder is populated by GNG prototypes. In the tested construction this is approximately two fixed anchors, 599 guards, and 199 prototypes. Duplicate filtering can shift the category count, but the final node cap remains identical. Guarded and pure GNG inspect a 4000-candidate pool, whereas Halton/PRM accepts only enough candidates to fill its graph; the methods are node-budget matched, not sample- or compute-budget matched. The 75% fraction was selected only on a development set and was frozen before held-out evaluation."
    )
    doc.add_heading("C. Cached Broad Phase and Exact Replanning", level=2)
    doc.add_paragraph(
        "Each roadmap node and interpolated edge stores capsules spanning the moving arm, end effector, fingers, and camera payload. After an environment update, point-to-segment and segment-to-segment distances mark blocked nodes and edges without rebuilding the roadmap."
    )
    add_equation(doc, "blocked(c,o) iff distance(axis(c),o) < radius(c)+clearance", 3)
    doc.add_paragraph(
        "The surviving route is checked against robot meshes and environment spheres/cylinders in a MoveIt PlanningScene using FCL. States are interpolated at normalized-joint step 0.05. If an exact check rejects an edge, that edge is disabled for the current environment and Dijkstra search repeats, up to twenty times. This is a discrete obstacle-update planner: it does not predict continuous obstacle motion."
    )

    doc.add_heading("V. Implementation and Experimental Design", level=1)
    doc.add_paragraph(
        "The C++ implementation runs on ROS 2 Humble on the robot's AGX computer. MoveIt RobotModel/RobotState provide kinematics and collision state; PlanningScene and FCL validate detailed geometry. Each benchmark process ran in an isolated ROS domain, published a synthetic joint state, and never launched MoveGroup, ros2_control, the camera, a controller publisher, or a trajectory action client."
    )
    table_i_caption = add_caption(doc, "Table I. Fixed roadmap and query parameters.")
    keep_with_next(table_i_caption)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [3000, 1860, 4500]
    set_table_geometry(table, widths)
    for idx, value in enumerate(("Parameter", "Value", "Role")):
        shade_cell(table.rows[0].cells[idx], LIGHT_BLUE)
        add_table_text(table.rows[0].cells[idx], value, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    parameter_rows = [
        ("Roadmap nodes", "800", "Equal final node budget for all methods"),
        ("GNG training signals", "4000", "Valid normalized configurations"),
        ("Guard fraction", "0.75", "Frozen after development ablation"),
        ("Neighbour candidates", "10", "Shared k-NN connection stage"),
        ("Max normalized q distance", "0.75", "Reject long local connections"),
        ("Max Cartesian edge length", "0.14 m", "Restrict end-effector displacement"),
        ("Target intersection radius", "0.05 m", "Environment-to-robot overlap"),
        ("Obstacle clearance", "0.035 m", "Capsule broad-phase margin"),
        ("Body / exact step", "0.08 / 0.05", "Cached sweep / exact interpolation"),
        ("Environment radii", "0.012 / 0.006 m", "Node sphere / edge cylinder"),
        ("Maximum exact replans", "20", "Bound edge removal and re-search"),
    ]
    for values in parameter_rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            add_table_text(cells[idx], value, align=WD_ALIGN_PARAGRAPH.CENTER if idx == 1 else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(table, widths)

    doc.add_heading("A. Development and Held-Out Protocol", level=2)
    doc.add_paragraph(
        "A matched development benchmark on deterministic offsets 0–49 exposed eight obstacle-update failures for pure GNG (offsets 6, 9, 16, 18, 35, 37, 44, and 47). Start and target were privileged anchor nodes inserted before sampling. Geometric diagnosis showed that the anchored home target node was always removed by the synthetic midpoint point obstacle: its minimum capsule margin was -6.04 mm, whereas the start state was only +0.119 mm outside the broad-phase threshold. Success therefore required an alternative node within the 50 mm target-intersection ball. The selected failure set was used only to choose the guard fraction; it was not reused for held-out claims."
    )
    ablation_caption = add_caption(doc, "Table II. Development-only guard-fraction ablation on the eight selected pure-GNG failure offsets.")
    keep_with_next(ablation_caption)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [2800, 2200, 4360]
    set_table_geometry(table, widths)
    for idx, value in enumerate(("Construction", "Success", "Interpretation")):
        shade_cell(table.rows[0].cells[idx], LIGHT_GRAY)
        add_table_text(table.rows[0].cells[idx], value, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    ablation_rows = [
        ("Pure GNG (0%)", "0/8", "Development failure set by definition"),
        ("Guarded GNG (10%)", "1/8", "Insufficient local support"),
        ("Guarded GNG (25%)", "4/8", "Partial recovery"),
        ("Guarded GNG (50%)", "7/8", "One remaining failure"),
        ("Guarded GNG (75%)", "8/8", "Selected and then frozen"),
    ]
    for values in ablation_rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            add_table_text(cells[idx], value, align=WD_ALIGN_PARAGRAPH.CENTER if idx == 1 else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(table, widths)
    doc.add_paragraph(
        "Held-out offsets 50–99 then generated 50 matched blocks, each containing all three methods with the same anchored target, fixed midpoint obstacle, Halton offset, node cap, and validity rules. All six method orders were cycled, but 50 is not divisible by six, leaving a small position imbalance. This produced 150 roadmap builds, 150 clear-scene queries, and 150 obstacle-update queries. Clear-scene success was 150/150. The result and exact-valid flags agreed on every query; exact-valid is the reported success criterion."
    )

    doc.add_heading("B. Statistical Analysis", level=2)
    doc.add_paragraph(
        "Obstacle-update success was the primary outcome. We report Wilson binomial intervals, Cochran's Q across the three paired methods, paired-bootstrap risk-difference intervals, and exact McNemar tests for the two prespecified guarded-GNG comparisons with Holm correction. Continuous paired outcomes use bootstrap intervals and Wilcoxon signed-rank tests with Holm correction within each metric; rank-biserial correlation reports paired effect direction. Timing after obstacle update is compared only on jointly successful offset pairs to avoid changing the analysis population by method. Medians and interquartile ranges are primary descriptive summaries because timing distributions are skewed. An exact sign test was used as a sensitivity check where Wilcoxon symmetry was questionable."
    )

    doc.add_heading("VI. Held-Out Results", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(results), width=Inches(6.35))
    doc_pr = run._r.xpath(".//wp:docPr")
    if doc_pr:
        doc_pr[0].set("descr", "Four-panel chart of held-out success, components, build time, and successful obstacle-update planning time")
    add_caption(doc, "Fig. 2. Development selection and held-out fixed-task results. The left panel is conditional on eight selected GNG failures and is not confirmatory. Held-out panels summarize deterministic offsets 50–99; success intervals and paired contrasts appear in Tables III–IV.")

    summary_caption = add_caption(doc, "Table III. Held-out benchmark across 50 matched deterministic sample-stream offsets. Continuous outcomes are median [IQR].")
    keep_with_next(summary_caption)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [3500, 1950, 1950, 1960]
    set_table_geometry(table, widths)
    for idx, value in enumerate(("Metric", "GNG", "Guarded GNG", "Halton/PRM")):
        shade_cell(table.rows[0].cells[idx], LIGHT_BLUE)
        add_table_text(table.rows[0].cells[idx], value, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    summary_rows = [
        ("Obstacle-update success", "46/50 (92%)", "49/50 (98%)", "49/50 (98%)"),
        ("Wilson 95% CI", "81.2–96.8%", "89.5–99.6%", "89.5–99.6%"),
        ("Connected components", "1 [1, 2]", "12 [10, 13]", "12 [10, 14]"),
        ("Validated edges", "4541 [4516, 4567]", "3627 [3571, 3667]", "3257 [3233, 3278]"),
        ("Build time (ms)", "1735.14 [1715.49, 1745.79]", "1625.67 [1604.22, 1645.90]", "1402.22 [1382.43, 1418.06]"),
        ("Clear query (ms)", "3.088 [2.892, 3.341]", "2.936 [2.761, 3.234]", "3.348 [3.006, 4.615]"),
        ("Obstacle query (ms)*", "19.233 [16.036, 25.165]", "18.556 [15.748, 27.415]", "26.554 [17.829, 37.446]"),
        ("Successful runs with >=1 replan", "1/46", "4/49", "7/49"),
        ("Total exact replans", "2", "8", "12"),
    ]
    for values in summary_rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            add_table_text(cells[idx], value, align=WD_ALIGN_PARAGRAPH.CENTER if idx else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(table, widths)
    p = doc.add_paragraph("*Each method's successful queries only (n=46, 49, and 49); these marginal summaries are not directly comparable. Pairwise timing inference uses only jointly successful offsets.")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.runs[0], size=8.5, italic=True, color=MUTED)

    effects_caption = add_caption(doc, "Table IV. Prespecified guarded-GNG paired effects on held-out offsets. Differences are guarded minus comparator with paired-bootstrap 95% CIs; success uses exact McNemar Holm correction across the two primary contrasts, while continuous Wilcoxon p-values are Holm-adjusted within metric.")
    keep_with_next(effects_caption)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    widths = [2600, 1600, 3000, 720, 1440]
    set_table_geometry(table, widths)
    for idx, value in enumerate(("Outcome", "Comparator", "Guarded minus comparator (95% CI)", "n", "Holm p")):
        shade_cell(table.rows[0].cells[idx], LIGHT_GRAY)
        add_table_text(table.rows[0].cells[idx], value, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE, size=8.8)
    set_repeat_table_header(table.rows[0])
    effect_rows = [
        ("Success risk difference", "GNG", "+6 pp [-2, 14]", "50", ".750"),
        ("Success risk difference", "Halton", "0 pp [-6, 6]", "50", "1.000"),
        ("Components", "GNG", "+10.60 [9.76, 11.48]", "50", "<.001"),
        ("Components", "Halton", "-0.06 [-1.28, 1.18]", "50", ".768"),
        ("Build time (ms)", "GNG", "-108.18 [-129.29, -87.60]", "50", "<.001"),
        ("Build time (ms)", "Halton", "+224.43 [206.08, 242.70]", "50", "<.001"),
        ("Clear query (ms)", "GNG", "-0.267 [-0.551, -0.011]", "50", ".091"),
        ("Clear query (ms)", "Halton", "-0.732 [-1.035, -0.446]", "50", "<.001"),
        ("Obstacle query (ms)", "GNG", "+2.734 [-0.716, 6.717]", "45", ".134"),
        ("Obstacle query (ms)", "Halton", "-3.626 [-7.228, 0.148]", "48", ".125"),
    ]
    for values in effect_rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            add_table_text(cells[idx], value, align=WD_ALIGN_PARAGRAPH.CENTER if idx else WD_ALIGN_PARAGRAPH.LEFT, size=8.8)
    set_table_geometry(table, widths)

    doc.add_heading("A. Primary Outcome", level=2)
    doc.add_paragraph(
        "The global success test did not reject equality across methods (Cochran Q(2)=3.00, p=.223). Guarded GNG succeeded on three more offsets than pure GNG, with discordant outcomes 4:1, but the paired interval included both a small loss and a meaningful gain. Against Halton/PRM, discordances were 1:1. The data therefore support feasibility and a promising point estimate, not superiority. Failures were target_intersection_blocked_or_disconnected: offsets 70, 76, 93, and 95 for GNG; offset 69 for guarded GNG; and offset 96 for Halton/PRM."
    )
    doc.add_heading("B. Connectivity, Cost, and Lazy Exact Validation", level=2)
    doc.add_paragraph(
        "Pure GNG formed approximately one to two large components and 1285 more edges than Halton/PRM on average. Guarding increased raw low-discrepancy support but raised the component count by 10.60 relative to GNG, making it statistically indistinguishable from Halton/PRM on this measure. Guarded GNG built 108 ms faster than pure GNG but 224 ms slower than Halton/PRM. Thus the 75% guard behaves as a hybrid tilted strongly toward local sample support and relinquishes pure GNG's low-component-count characteristic."
    )
    doc.add_paragraph(
        "The exact-rejection branch was exercised: pure GNG required exact replanning in one held-out offset (two rejected edges), guarded GNG in four offsets (eight events), and Halton/PRM in seven offsets (twelve events). These paths ultimately validated, showing that capsule filtering and exact mesh checking are not redundant. Obstacle-update timing differences for guarded GNG were inconclusive against both comparators on jointly successful offsets."
    )

    doc.add_heading("VII. Discussion and Limitations", level=1)
    doc.add_paragraph(
        "The central result is that global component count and obstacle-conditioned local target support are different objectives. GNG prototypes efficiently connect the dominant valid-configuration distribution, yet may underrepresent a small target neighbourhood after the nominal target branch is pruned. Deterministic guards recover local alternatives but fragment the graph under the present shared edge thresholds. Because workspace coverage and dispersion were not measured, broader coverage claims remain a hypothesis. A stronger next method should preserve selected GNG adjacency, add multiple target anchors at query time, or repair disconnected guard nodes toward the start component rather than increasing the guard fraction globally."
    )
    limitations = [
        ("Scene generalization", "All held-out offsets share one anchored start, target, midpoint point obstacle, and clearance geometry. Offsets measure construction variability, not generalization across tasks or scenes."),
        ("Development selection", "The 75% fraction was chosen on eight observed GNG failures. It is cleanly separated from offsets 50–99, but additional held-out scenes are needed before fixing it as a general default."),
        ("Budget fairness", "Methods have equal nodes and shared local rules, but not equal candidate information, edge counts, build time, or collision-check workload. Future comparisons should include node-, edge-, candidate-, and wall-clock-matched budgets."),
        ("Stress geometry", "The anchored home target is always broad-phase blocked, making the query intentionally dependent on a nearby alternative. A task matrix must include clear, narrow-passage, edge-only, and target-adjacent obstacles."),
        ("Perception", "The benchmark uses synthetic graph primitives and does not measure depth noise, semantic error, DD-GNG update delay, or transform uncertainty."),
        ("Baselines", "RRTConnect, PRM*, SPARS, and topology-repair variants are not yet included under matched collision and time budgets."),
        ("Timing", "One clear and one obstacle-update query were timed per build under normal AGX background load; method-position counts were slightly imbalanced. Repeated within-build timing and controlled load are needed."),
        ("Implementation fidelity", "Pure GNG trains 800 prototypes before two are truncated by fixed anchors, and the runner records requested rather than echoed runtime parameters. Correct these issues and rerun before submission."),
        ("Deployment default", "The evaluated 75%-guarded configuration is explicit benchmark input; the package default remains pure GNG with a different guard fraction. Deployment claims must state the effective runtime parameters."),
        ("Provenance", "The benchmark source is hashed but remains based on a dirty parent worktree. Archive a commit/tag plus URDF, SRDF, MoveIt/FCL, compiler, and dependency versions."),
        ("Execution and safety", "No physical trajectory was sent. Velocity, acceleration, torque, controller tracking, calibration error, emergency-stop procedure, and human-safe trials remain unevaluated."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    widths = [2150, 7210]
    set_table_geometry(table, widths)
    for idx, value in enumerate(("Limitation", "Implication / next experiment")):
        shade_cell(table.rows[0].cells[idx], LIGHT_GRAY)
        add_table_text(table.rows[0].cells[idx], value, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for label, detail in limitations:
        cells = table.add_row().cells
        add_table_text(cells[0], label, bold=True)
        add_table_text(cells[1], detail)
    set_table_geometry(table, widths)
    add_callout(
        doc,
        "Immediate ICRA-critical next experiment",
        "Freeze the algorithms and evaluate a preregistered matrix of unseen start-target pairs and obstacle scenes. Use independent scrambled streams with paired offsets within each scene, include RRTConnect and a sparse-roadmap baseline, report success as the primary endpoint, and separate construction, invalidation, graph-search, and exact-validation time. Only then add controlled physical executions.",
    )

    doc.add_heading("VIII. Conclusion", level=1)
    doc.add_paragraph(
        "This work presents a dual-topology, preview-only reachability planner with configuration-retaining robot nodes, semantic environment targets, cached capsule invalidation, and lazy MoveIt/FCL edge rejection. Across 50 paired held-out deterministic offsets, a 75%-guarded GNG construction raised the obstacle-update success point estimate from 92% to 98% and matched Halton/PRM, but neither paired comparison was statistically significant. Pure GNG retained a much lower component count, while guarded GNG became structurally similar to the low-discrepancy baseline. The defensible fixed-task conclusion is a target-robustness–connectivity hypothesis, not general coverage or method superiority. The next research step is not more tuning on these offsets, but topology-aware repair and evaluation on unseen multi-scene tasks under matched computational budgets."
    )

    doc.add_heading("References", level=1)
    references = [
        "B. Fritzke, 'A Growing Neural Gas Network Learns Topologies,' in Advances in Neural Information Processing Systems 7, pp. 625–632, 1995.",
        "L. E. Kavraki, P. Svestka, J.-C. Latombe, and M. H. Overmars, 'Probabilistic Roadmaps for Path Planning in High-Dimensional Configuration Spaces,' IEEE Transactions on Robotics and Automation, vol. 12, no. 4, pp. 566–580, 1996, doi: 10.1109/70.508439.",
        "M. S. Branicky, S. M. LaValle, K. Olson, and L. Yang, 'Quasi-Randomized Path Planning,' in Proc. IEEE International Conference on Robotics and Automation, pp. 1481–1487, 2001, doi: 10.1109/ROBOT.2001.932820.",
        "F. Zacharias, C. Borst, and G. Hirzinger, 'Capturing Robot Workspace Structure: Representing Robot Capabilities,' in Proc. IEEE/RSJ International Conference on Intelligent Robots and Systems, pp. 3229–3236, 2007, doi: 10.1109/IROS.2007.4399105.",
        "O. Porges, T. Stouraitis, C. Borst, and M. A. Roa, 'Reachability and Capability Analysis for Manipulation Tasks,' in ROBOT2013: First Iberian Robotics Conference, pp. 703–718, 2014, doi: 10.1007/978-3-319-03653-3_50.",
        "R. Bohlin and L. E. Kavraki, 'Path Planning Using Lazy PRM,' in Proc. IEEE International Conference on Robotics and Automation, pp. 521–528, 2000, doi: 10.1109/ROBOT.2000.844107.",
        "M. Kallmann and M. J. Mataric, 'Motion Planning Using Dynamic Roadmaps,' in Proc. IEEE International Conference on Robotics and Automation, pp. 4399–4404, 2004, doi: 10.1109/ROBOT.2004.1302410.",
        "L. Jaillet and T. Simeon, 'A PRM-Based Motion Planner for Dynamically Changing Environments,' in Proc. IEEE/RSJ International Conference on Intelligent Robots and Systems, pp. 1606–1611, 2004, doi: 10.1109/IROS.2004.1389625.",
        "A. Dobson and K. E. Bekris, 'Sparse Roadmap Spanners for Asymptotically Near-Optimal Motion Planning,' International Journal of Robotics Research, vol. 33, no. 1, pp. 18–47, 2014, doi: 10.1177/0278364913498292.",
        "M. Phillips, B. Cohen, S. Chitta, and M. Likhachev, 'E-Graphs: Bootstrapping Planning with Experience Graphs,' in Robotics: Science and Systems VIII, 2012, doi: 10.15607/RSS.2012.VIII.043.",
        "D. Coleman, I. A. Sucan, M. Moll, K. Okada, and N. Correll, 'Experience-Based Planning with Sparse Roadmap Spanners,' in Proc. IEEE International Conference on Robotics and Automation, pp. 900–905, 2015, doi: 10.1109/ICRA.2015.7139284.",
        "B. Ichter, J. Harrison, and M. Pavone, 'Learning Sampling Distributions for Robot Motion Planning,' in Proc. IEEE International Conference on Robotics and Automation, pp. 7087–7094, 2018, doi: 10.1109/ICRA.2018.8460730.",
        "A. H. Qureshi, A. Simeonov, M. J. Bency, and M. C. Yip, 'Motion Planning Networks,' in Proc. IEEE International Conference on Robotics and Automation, pp. 2118–2124, 2019, doi: 10.1109/ICRA.2019.8793889.",
        "C. Chamzas, Z. Kingston, C. Quintero-Pena, A. Shrivastava, and L. E. Kavraki, 'Learning Sampling Distributions Using Local 3D Workspace Decompositions for Motion Planning in High Dimensions,' in Proc. IEEE International Conference on Robotics and Automation, pp. 1283–1289, 2021, doi: 10.1109/ICRA48506.2021.9561104.",
        "D. Coleman, I. A. Sucan, S. Chitta, and N. Correll, 'Reducing the Barrier to Entry of Complex Robotic Software: a MoveIt! Case Study,' Journal of Software Engineering for Robotics, vol. 5, no. 1, pp. 3–16, 2014, doi: 10.6092/JOSER_2014_05_01_p3.",
        "J. Pan, S. Chitta, and D. Manocha, 'FCL: A General Purpose Library for Collision and Proximity Queries,' in Proc. IEEE International Conference on Robotics and Automation, 2012, doi: 10.1109/ICRA.2012.6225337.",
    ]
    for idx, reference in enumerate(references, 1):
        add_reference(doc, idx, reference)

    doc.core_properties.title = "Target-Robustness–Connectivity Trade-offs in GNG-Quantized Reachability Roadmaps"
    doc.core_properties.subject = "Held-out 50-offset ICRA-oriented working manuscript"
    doc.core_properties.keywords = "GNG, reachability roadmap, Halton, manipulation, collision checking, ROS 2"
    doc.core_properties.comments = "Working draft with development/held-out separation; not yet a multi-scene or physical-robot study."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
